from docx import Document

doc = Document()
doc.add_heading('RESOLUCIÓN DE INTENDENCIA N° 00{{ numExpediente }}-2024', 1)
doc.add_paragraph('Razón Social: {{ razonSocial }}')
doc.add_paragraph('RUC: {{ ruc }}')
doc.add_paragraph('Domicilio: {{ domicilio }}')
doc.add_paragraph('')
doc.add_paragraph('Por la presente, en relación a la boleta {{ boleta }} con un monto de {{ monto }}, se procede a evaluar su expediente.')
doc.add_paragraph('')
doc.add_paragraph('Atentamente,')
doc.add_paragraph('{{ firma }}')
doc.add_paragraph('{{ imagen_firma }}')

doc.save('app/modulos/autodoc/plantillas/Plantilla_Base.docx')
print("Plantilla_Base.docx actualizada con soportes para firmas.")
