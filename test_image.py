from docxtpl import DocxTemplate, InlineImage
from docx.shared import Mm
import tempfile
import os

def test_image_injection():
    # 1. Crear una plantilla de prueba con una imagen inyectada
    from docx import Document
    doc = Document()
    doc.add_paragraph('Ejemplo de inyección de firma:')
    doc.add_paragraph('{{ firma_imagen }}')
    doc.save('test_img_template.docx')
    
    # 2. Crear una imagen de prueba temporal (un cuadrado azul simulando un sello)
    from PIL import Image
    img = Image.new('RGB', (200, 100), color = (0, 102, 179))
    img.save('firma_prueba.png')
    
    # 3. Probar inyección con docxtpl
    tpl = DocxTemplate('test_img_template.docx')
    
    # Aquí es donde le decimos a docxtpl que reemplace {{ firma_imagen }} por una InlineImage
    # Mm(50) establece el ancho de la imagen en milímetros para que no pierda su lugar
    imagen_sello = InlineImage(tpl, 'firma_prueba.png', width=Mm(50))
    
    context = {
        'firma_imagen': imagen_sello
    }
    
    tpl.render(context)
    tpl.save('resultado_firma.docx')
    print("Prueba de imagen exitosa. Revisa resultado_firma.docx")

if __name__ == '__main__':
    test_image_injection()
